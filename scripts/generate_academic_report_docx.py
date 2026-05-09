"""
Build a long-form academic DOCX report for the Bihar groundwater–HMIS project.

Usage (from repository root):
    python scripts/generate_academic_report_docx.py

Output:
    results/reports/Academic_Project_Report_Full.docx

Requires: python-docx
"""

from __future__ import annotations

import json
import sys
from pathlib import Path
from typing import Iterable

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "scripts"))

from docx import Document  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING  # type: ignore
from docx.shared import Inches, Pt, RGBColor  # type: ignore

import academic_report_body as rb  # noqa: E402
import academic_report_extended as ext  # noqa: E402


METRICS_PATH = ROOT / "results" / "models" / "metrics.json"
SHAP_PATH = ROOT / "results" / "models" / "shap_importance.csv"
FIG_DIR = ROOT / "results" / "figures"
OUT_DOCX = ROOT / "results" / "reports" / "Academic_Project_Report_Full.docx"


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for lvl in range(1, 4):
        st = doc.styles[f"Heading {lvl}"]
        st.font.name = "Times New Roman"
        st.font.color.rgb = RGBColor(0x1A, 0x24, 0x5E)
        if lvl == 1:
            st.font.size = Pt(16)
            st.font.bold = True
        elif lvl == 2:
            st.font.size = Pt(14)
            st.font.bold = True
        else:
            st.font.size = Pt(12)
            st.font.bold = True


def add_title_block(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x4A)

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(12)
    r2.italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run(
        "Academic submission manuscript generated from the reproducible code repository.\n"
        "Institutional affiliation and author names should be inserted here before final deposit."
    )
    m.font.size = Pt(11)

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, *, first_indent: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(0.25)
    p.add_run(text)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_table_grid(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri + 1].cells[ci].text = val


def add_figure(doc: Document, path: Path, caption: str, width_in: float = 6.2) -> None:
    if path.exists():
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width_in))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(10)
        doc.add_paragraph()
    else:
        doc.add_paragraph(f"[Figure missing: {path.name} — run the pipeline to generate it.]")
        doc.add_paragraph()


def load_metrics() -> dict:
    if not METRICS_PATH.exists():
        return {}
    return json.loads(METRICS_PATH.read_text(encoding="utf-8"))


def expanded_ion_methodology() -> list[str]:
    """Additional pedagogical depth for page count and teaching clarity."""
    return split_blocks(
        """
Each ion tracked in the merged dataset corresponds to a measurable chemical constituent frequently monitored in national groundwater programmes because of relevance to taste, scaling, corrosion of pipes, irrigation suitability, or human-health guidelines under specific exposure regimes. pH influences solubility of metals and speciation of contaminants; although extreme pH alone does not diagnose toxicity, shifts often co-occur with mineralisation gradients worth describing to hydrogeologists collaborating with health teams.

Electrical conductivity offers an aggregate correlate of ionic strength; districts rising in conductivity alongside chloride and sodium-type patterns may indicate salinity ingress in coastal-influenced aquifers or evaporative concentration in inland traps, though definitive hydrogeological typing exceeds this pipeline’s scope. Bicarbonate derives from carbonate equilibrium and microbial processes in some contexts; chloride and sulphate often participate in identifying anthropogenic inputs or evaporite dissolution depending on local geology.

Nitrate attracts attention because elevated concentrations in drinking water have been linked epidemiologically to infant methemoglobinemia in settings where contaminated wells supply formula-fed infants; ecological associations at district scale cannot substantiate individual risk yet justify prioritising surveillance alignment between chemistry labs and maternal-child health programmes.

Fluoride monitoring addresses endemic fluorosis concerns documented historically across parts of India; district medians smooth extreme hotspots but still orient further spatial refinement.

Total dissolved solids integrate multiple ions into a single mineralisation indicator comparable across laboratories when methodology is consistent.

Monitoring intensity captured through well counts acknowledges that districts rarely sampled may yield median chemistry estimates with higher uncertainty; treating counts as auxiliary context avoids pretending equal confidence everywhere.

Transform choices — medians, clipping negatives for non-negative physical quantities, median imputation inside folds — reflect pragmatic hygiene oriented toward stability and auditability rather than maximal exploitation of distributional modelling.
"""
    )


def split_blocks(big_text: str) -> list[str]:
    return [p.strip() for p in big_text.strip().split("\n\n") if p.strip()]


def results_section_text(m: dict) -> list[str]:
    if not m:
        return split_blocks(
            """
Results artefacts were not found at the expected metrics path when this document was generated. Run `python scripts/run_pipeline.py` after building `data/processed/official_district_year_master.csv` to populate `results/models/metrics.json` and figures, then regenerate this report. Placeholder text cannot substitute for empirical outputs in a final submission.
"""
        )

    po = m.get("pooled_oof_metrics_best_model", {})
    lines = [
        f"The fitted configuration summarised here reflects the saved pipeline run embedded in the repository artefact `metrics.json`. "
        f"The analytical sample contained {m.get('n_rows', 'N/A')} district-year rows spanning {m.get('n_districts', 'N/A')} distinct districts. "
        f"Grouped cross-validation used {m.get('spatial_cv_splits', 'N/A')} folds constrained by district identifiers.",
        "Table R1 lists mean fold metrics for each candidate algorithm. Lower root mean squared error indicates better average fit across held-out districts, though interpretation must remain cautious because ecological bias and spatial non-stationarity may dominate.",
        "The pooled out-of-fold summary stitches predictions from each validation fold into one vector aligned with all rows, then recomputes global error metrics. This summary often reads differently from per-fold R-squared averages when folds contain different numbers of heterogeneous districts.",
        f"For the selected best model (`{m.get('best_model', '?')}`), pooled out-of-fold root mean squared error equalled {po.get('rmse', float('nan')):.4f} on the stabilised log-burden scale. "
        f"Pooled R-squared equalled {po.get('r2', float('nan')):.4f}. Spearman rank correlation between pooled predictions and observed outcomes was {po.get('spearman_r', float('nan')):.4f} "
        f"(two-sided p-value {po.get('spearman_p_two_sided', float('nan')):.4g}). "
        f"The naive mean baseline achieved root mean squared error {po.get('naive_mean_baseline_rmse', float('nan')):.4f}, "
        f"and the relative percentage improvement of the model versus that baseline was {po.get('rmse_improvement_vs_naive_pct', float('nan')):.2f} percent (negative values mean the model was worse than predicting the global mean).",
        "These numerical outcomes reinforce the report’s consistent messaging: strict spatial validation combined with a narrow environmental feature set yields limited predictive power for the chosen HMIS composite. Negative improvements versus naive baselines should be presented frankly in academic defence as evidence that groundwater summaries alone are insufficient anchors for forecasting burden ranks without richer contextual covariates.",
        "Despite modest discrimination, diagnostic plots still reveal systematic tendencies useful for critical reflection — for example clustering of errors by districts that repeatedly appear in high-residual regions, suggesting unmodelled structural factors such as localised floods or reporting bursts.",
    ]
    return lines


def main() -> None:
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    add_title_block(doc, rb.TITLE, rb.SUBTITLE)

    add_heading(doc, "Table of Contents", 1)
    toc_lines = [
        "Abstract",
        "Keywords",
        "1. Introduction",
        "2. Literature Review",
        "3. Objectives",
        "4. Study Area and Analytical Scale",
        "5. Data Sources and Provenance",
        "6. Dataset Construction and Row Logic",
        "7. Variable Definitions and Scientific Interpretation",
        "7.1 Expanded discussion of groundwater chemistry parameters",
        "8. Methodology",
        "8.1 Extract–transform–load and harmonisation",
        "8.2 Machine learning protocol",
        "8.3 Susceptibility and environmental stress profiling",
        "8.4 Ethics, governance, and communication of uncertainty",
        "9. Algorithms Used — Detailed Account",
        "9.1 Ridge regression",
        "9.2 Random forest",
        "9.3 Extreme gradient boosting",
        "9.4 Validation metrics and explainability integration",
        "10. Results and Analysis",
        "10.1 Statistical primer: metrics, spatial validation, and cautious reading",
        "10.2 Policy context and programme communication",
        "10.3 Worked example: linking one district-year storyline",
        "10.4 Integrated synthesis across data, methods, and interpretation",
        "11. Discussion",
        "12. Limitations",
        "13. Conclusion",
        "14. Future Scope",
        "15. Appendices (A–G)",
        "References",
    ]
    for line in toc_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Inches(0.2)
    doc.add_page_break()

    add_heading(doc, "Abstract", 1)
    for para in rb.ABSTRACT:
        add_para(doc, para)

    add_heading(doc, "Keywords", 1)
    add_para(doc, rb.KEYWORDS)

    doc.add_page_break()

    add_heading(doc, "1. Introduction", 1)
    for para in rb.INTRODUCTION:
        add_para(doc, para)

    doc.add_page_break()

    add_heading(doc, "2. Literature Review", 1)
    for para in rb.LITERATURE_REVIEW:
        add_para(doc, para)
    for para in rb.EXTRA_LITERATURE_BRIDGE:
        add_para(doc, para)

    doc.add_page_break()

    add_heading(doc, "3. Objectives", 1)
    for para in rb.OBJECTIVES:
        add_para(doc, para)

    doc.add_page_break()

    add_heading(doc, "4. Study Area and Analytical Scale", 1)
    for para in rb.STUDY_AREA:
        add_para(doc, para)

    doc.add_page_break()

    add_heading(doc, "5. Data Sources and Provenance", 1)
    for para in rb.DATA_SOURCES:
        add_para(doc, para)

    doc.add_page_break()

    add_heading(doc, "6. Dataset Construction and Row Logic", 1)
    for para in rb.DATASET_DESCRIPTION:
        add_para(doc, para)

    add_heading(doc, "Table 1 — Core modelling columns (conceptual schema)", 2)
    add_table_grid(
        doc,
        ["Column group", "Representative fields", "Role in analysis"],
        [
            [
                "Identifiers",
                "District, Year",
                "Merge keys, grouping for spatial folds, time alignment.",
            ],
            [
                "Groundwater chemistry",
                "GW_pH_median, GW_EC_uScm_median, GW_HCO3_median, GW_Cl_median, GW_SO4_median, GW_NO3_median, GW_F_mgL_median, GW_TDS_median",
                "Primary predictors summarising district-year water quality.",
            ],
            [
                "Monitoring",
                "n_wells",
                "Intensity of observations feeding medians; contextual uncertainty proxy.",
            ],
            [
                "Health burden",
                "HMIS_burden_raw, HMIS_vulnerability_z",
                "Composite burden sum and multi-indicator z summary for description.",
            ],
            [
                "Demography",
                "Population_2011",
                "Denominator for per-capita rates.",
            ],
            [
                "Modelling target",
                "HMIS_burden_per_100k, Target_log_burden_per100k",
                "Rate and stabilised log-rate response.",
            ],
            [
                "Stratification aid",
                "High_risk_tercile",
                "Descriptive tertiles when quantiles succeed.",
            ],
        ],
    )
    doc.add_paragraph()

    doc.add_page_break()

    add_heading(doc, "7. Variable Definitions and Scientific Interpretation", 1)
    for para in rb.VARIABLE_EXPLANATIONS:
        add_para(doc, para)

    add_heading(doc, "7.1 Expanded discussion of groundwater chemistry parameters", 2)
    for para in expanded_ion_methodology():
        add_para(doc, para)

    doc.add_page_break()

    add_heading(doc, "8. Methodology", 1)
    add_heading(doc, "8.1 Extract–transform–load and harmonisation", 2)
    for para in rb.METHODOLOGY_ETL:
        add_para(doc, para)

    add_heading(doc, "8.2 Machine learning protocol", 2)
    for para in rb.METHODOLOGY_ML:
        add_para(doc, para)

    add_heading(doc, "8.3 Susceptibility and environmental stress profiling", 2)
    for para in rb.METHODOLOGY_SUSCEPTIBILITY:
        add_para(doc, para)

    add_heading(doc, "8.4 Ethics, governance, and communication of uncertainty", 2)
    for para in rb.METHODOLOGY_ETHICS:
        add_para(doc, para)

    doc.add_page_break()

    add_heading(doc, "9. Algorithms Used — Detailed Account", 1)
    add_heading(doc, "9.1 Ridge regression", 2)
    for para in rb.ALGORITHM_RIDGE:
        add_para(doc, para)

    add_heading(doc, "9.2 Random forest", 2)
    for para in rb.ALGORITHM_RF:
        add_para(doc, para)

    add_heading(doc, "9.3 Extreme gradient boosting", 2)
    for para in rb.ALGORITHM_XGB:
        add_para(doc, para)

    add_heading(doc, "9.4 Validation metrics and explainability integration", 2)
    for para in rb.METRICS_AND_VALIDATION:
        add_para(doc, para)

    doc.add_page_break()

    add_heading(doc, "10. Results and Analysis", 1)
    m = load_metrics()
    for para in results_section_text(m):
        add_para(doc, para)

    add_heading(doc, "10.1 Statistical primer: metrics, spatial validation, and cautious reading", 2)
    for para in ext.STATISTICAL_AND_VALIDATION_PRIMER:
        add_para(doc, para)

    add_heading(doc, "Table R1 — Mean spatial cross-validation metrics by model", 2)
    rows = []
    if m:
        for item in m.get("cv_by_model", []):
            rows.append(
                [
                    str(item.get("model", "")),
                    f"{item.get('rmse_mean', float('nan')):.4f}",
                    f"{item.get('mae_mean', float('nan')):.4f}",
                    f"{item.get('r2_mean', float('nan')):.4f}",
                ]
            )
    else:
        rows = [["—", "—", "—", "—"]]
    add_table_grid(doc, ["Model", "Mean RMSE", "Mean MAE", "Mean R²"], rows)
    doc.add_paragraph()

    add_heading(doc, "Table R2 — Pooled out-of-fold summary (best model)", 2)
    po = m.get("pooled_oof_metrics_best_model", {}) if m else {}
    add_table_grid(
        doc,
        ["Metric", "Value"],
        [
            ["Best model name", str(m.get("best_model", "n/a"))],
            ["Pooled out-of-fold RMSE", f"{po.get('rmse', float('nan')):.6f}"],
            ["Pooled out-of-fold R²", f"{po.get('r2', float('nan')):.6f}"],
            ["Spearman rho (pred vs observed)", f"{po.get('spearman_r', float('nan')):.6f}"],
            ["Spearman two-sided p-value", f"{po.get('spearman_p_two_sided', float('nan')):.6g}"],
            ["Naive mean baseline RMSE", f"{po.get('naive_mean_baseline_rmse', float('nan')):.6f}"],
            ["RMSE improvement vs naive (%)", f"{po.get('rmse_improvement_vs_naive_pct', float('nan')):.4f}"],
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "Table R3 — Mean absolute SHAP values (global importance)", 2)
    shap_rows: list[list[str]] = []
    if SHAP_PATH.exists():
        import csv

        with SHAP_PATH.open(newline="", encoding="utf-8") as f:
            r = csv.DictReader(f)
            for row in r:
                shap_rows.append([row["feature"], f"{float(row['mean_abs_shap']):.6f}"])
    if not shap_rows:
        shap_rows = [["(regenerate pipeline for SHAP CSV)", "—"]]
    add_table_grid(doc, ["Feature", "Mean |SHAP|"], shap_rows[:15])
    doc.add_paragraph()

    add_heading(doc, "Figures", 2)
    figs = [
        (FIG_DIR / "actual_vs_predicted.png", "Figure 1. Out-of-fold predictions versus observed stabilised log burden."),
        (FIG_DIR / "residuals.png", "Figure 2. Residual distribution for examining systematic bias."),
        (FIG_DIR / "shap_summary.png", "Figure 3. SHAP summary (mean absolute impact ranking)."),
        (FIG_DIR / "top_districts_pred.png", "Figure 4. Districts with relatively elevated model predictions — exploratory prioritisation view."),
    ]
    for fp, cap in figs:
        add_figure(doc, fp, cap)

    add_heading(doc, "10.2 Policy context and programme communication", 2)
    for para in ext.POLICY_AND_PUBLIC_HEALTH_CONTEXT:
        add_para(doc, para)

    add_heading(doc, "10.3 Worked example: linking one district-year storyline", 2)
    for para in ext.WORKED_EXAMPLE_INTERPRETATION:
        add_para(doc, para)

    add_heading(doc, "10.4 Integrated synthesis across data, methods, and interpretation", 2)
    for para in ext.INTEGRATED_SYNTHESIS_BEFORE_DISCUSSION:
        add_para(doc, para)

    doc.add_page_break()

    add_heading(doc, "11. Discussion", 1)
    for para in rb.DISCUSSION:
        add_para(doc, para)

    doc.add_page_break()

    add_heading(doc, "12. Limitations", 1)
    for para in rb.LIMITATIONS:
        add_para(doc, para)

    doc.add_page_break()

    add_heading(doc, "13. Conclusion", 1)
    for para in rb.CONCLUSION:
        add_para(doc, para)

    doc.add_page_break()

    add_heading(doc, "14. Future Scope", 1)
    for para in rb.FUTURE_SCOPE:
        add_para(doc, para)

    doc.add_page_break()

    add_heading(doc, "15. Appendices", 1)
    add_heading(doc, "Appendix A — Software stack and reproducibility", 2)
    for para in rb.APPENDIX_SOFTWARE:
        add_para(doc, para)
    add_heading(doc, "Appendix B — HMIS indicators used in the burden composite", 2)
    for para in rb.APPENDIX_HMIS_LIST:
        add_para(doc, para)
    add_heading(doc, "Appendix C — Command-line workflow (reference)", 2)
    for para in rb.APPENDIX_COMMANDS:
        add_para(doc, para)

    add_heading(doc, "Appendix D — Data quality review and assumption audit", 2)
    for para in ext.DATA_QUALITY_AND_ASSUMPTION_AUDIT:
        add_para(doc, para)

    add_heading(doc, "Appendix E — Replicator checklist and oral examination preparation", 2)
    for para in ext.REPLICATOR_CHECKLIST_AND_VIVA_GUIDE:
        add_para(doc, para)

    add_heading(doc, "Appendix F — Bihar context: groundwater, seasons, and health reporting", 2)
    for para in ext.CONTEXT_BIHAR_HYDROLOGY_HEALTH_SYSTEMS:
        add_para(doc, para)

    add_heading(doc, "Appendix G — Comparison of analytical alternatives and extension pathways", 2)
    for para in ext.COMPARISON_OF_ANALYTICAL_ALTERNATIVES:
        add_para(doc, para)

    doc.add_page_break()

    add_heading(doc, "References", 1)
    for ref in rb.REFERENCES_LIST:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.hanging_indent = Inches(-0.25)

    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
